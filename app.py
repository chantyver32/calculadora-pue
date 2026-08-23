import streamlit as st
import pandas as pd
from sqlalchemy import text
from datetime import datetime, timedelta
import pytz
import math
import urllib.parse
import io
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re  
import os
import streamlit.components.v1 as components

# ------------------ 1. CONFIGURACIÓN GENERAL ------------------
st.set_page_config(page_title="Insumos Champlitte", page_icon="⚖️", layout="wide")

with st.spinner('Iniciando sistema Champlitte... 🥐'):
    zona_mx = pytz.timezone('America/Mexico_City')
    fecha_hoy_mx = datetime.now(zona_mx).date()

st.markdown("""
    <style>
    .block-container { padding-top: 3rem; padding-bottom: 1rem; }
    .main { background-color: #f5f7f9; }
    
    .stButton > button, 
    .stFormSubmitButton > button { 
        width: 100%; border-radius: 8px; font-weight: bold; 
        transition: none !important; -webkit-transition: none !important;
    }
    .stButton > button:focus, .stButton > button:active,
    .stFormSubmitButton > button:focus, .stFormSubmitButton > button:active {
        box-shadow: none !important; outline: none !important; transform: none !important;
    }

    [data-testid="stElementContainer"], [data-testid="stForm"] {
        transition: none !important; animation: none !important;
    }

    /* CENTRADO DEL MENSAJE DE CONFIRMACIÓN (TOAST) */
    div[data-testid="stToastContainer"] {
        top: 2rem !important; 
        bottom: auto !important; 
        left: 50% !important; 
        right: auto !important; 
        transform: translateX(-50%) !important;
        justify-content: center !important;
        align-items: center !important;
    }
    
    ul[role="listbox"] li[aria-selected="true"] {
        background-color: transparent !important; font-weight: bold !important;
    }

    div[data-baseweb="popover"] > div {
        background-color: #1a1a1c !important; border-radius: 8px !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important; box-shadow: 0 4px 12px rgba(0,0,0,0.8) !important;
    }
    div[data-baseweb="popover"] ul { background-color: transparent !important; }
    div[data-baseweb="popover"] li {
        background-color: transparent !important; color: #FFFFFF !important; font-size: 14px !important; padding: 10px 0 !important;
    }
    div[data-baseweb="popover"] li:hover { background-color: #2d2d30 !important; }
    div[data-baseweb="popover"] li[aria-selected="true"] { background-color: #3a3b3e !important; font-weight: bold !important; }

    div[data-baseweb="select"] > div {
        background-color: #1a1a1c !important; border-radius: 8px !important; border: 1px solid rgba(255, 255, 255, 0.2) !important;
    }
    div[data-baseweb="select"] > div:focus-within { border-color: #ff4b4b !important; box-shadow: 0 0 0 1px #ff4b4b !important; }
    div[data-baseweb="select"] div, div[data-baseweb="select"] svg { color: #FFFFFF !important; fill: #FFFFFF !important; }

    .btn-wa {
        background-color: #25D366; color: white !important; padding: 10px 20px;
        text-align: center; text-decoration: none !important; display: block;
        font-size: 14px; font-weight: bold; border-radius: 8px; margin: 10px 0; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .btn-wa:hover { background-color: #128C7E; }
    
    div[data-testid="stMetricValue"] { font-size: 28px; color: #1f77b4; }
    div[data-testid="stMetricDelta"] { font-size: 30px !important; font-weight: bold !important; }
    </style>
""", unsafe_allow_html=True)

if "show_toast" in st.session_state:
    st.toast(st.session_state.show_toast)
    del st.session_state.show_toast

# ------------------ 2. CONEXIÓN A SUPABASE Y ORDENAMIENTO OFICIAL ------------------
db_url = os.environ.get("SUPABASE_URL")
if not db_url:
    try:
        db_url = st.secrets["SUPABASE_URL"]
    except (FileNotFoundError, KeyError):
        st.error("🚨 Error crítico: No se encontró 'SUPABASE_URL'.")
        st.stop() 

conn = st.connection("supabase", type="sql", url=db_url)

ORDEN_CATEGORIAS_OFICIAL = ["Papelería Venta", "Limpieza Venta", "Insumos Venta"]

with conn.session as s:
    s.execute(text('''CREATE TABLE IF NOT EXISTS pesajes_individuales 
                 (id SERIAL PRIMARY KEY, sucursal TEXT, fecha_hora TEXT, articulo TEXT, categoria TEXT, 
                 peso_bruto REAL, tara REAL, pue REAL, resultado_pue REAL, detalle_formula TEXT)'''))

    s.execute(text('''CREATE TABLE IF NOT EXISTS pesajes_guardados 
                 (id SERIAL PRIMARY KEY, sucursal TEXT, fecha_hora TEXT, articulo TEXT, categoria TEXT, 
                 peso_bruto REAL, tara REAL, pue REAL, resultado_pue REAL, detalle_formula TEXT)'''))

    s.execute(text('''CREATE TABLE IF NOT EXISTS auditoria_stock 
                 (id SERIAL PRIMARY KEY, sucursal TEXT, articulo TEXT, categoria TEXT, 
                 total_real REAL, stock REAL, diferencia REAL, UNIQUE(sucursal, articulo))'''))
    
    s.execute(text('''CREATE TABLE IF NOT EXISTS usuarios 
                 (id SERIAL PRIMARY KEY, username TEXT, password TEXT)'''))
                 
    s.execute(text('''CREATE TABLE IF NOT EXISTS catalogo_productos 
                 (id SERIAL PRIMARY KEY, categoria TEXT, articulo TEXT, pue REAL, UNIQUE(categoria, articulo))'''))

    s.execute(text('ALTER TABLE pesajes_individuales ADD COLUMN IF NOT EXISTS categoria TEXT;'))
    s.execute(text('ALTER TABLE pesajes_guardados ADD COLUMN IF NOT EXISTS categoria TEXT;'))
    s.execute(text('ALTER TABLE auditoria_stock ADD COLUMN IF NOT EXISTS categoria TEXT;'))
    s.execute(text('ALTER TABLE pesajes_guardados ADD COLUMN IF NOT EXISTS aplicado_en_corte BOOLEAN DEFAULT TRUE;'))
    s.execute(text('ALTER TABLE catalogo_productos ADD COLUMN IF NOT EXISTS ubicacion_conteo TEXT DEFAULT \'Combinado\';'))
    s.execute(text('ALTER TABLE catalogo_productos ADD COLUMN IF NOT EXISTS redondeo TEXT DEFAULT \'No\';'))

    s.commit()

# ------------------ DICCIONARIO BASE ------------------
dicc_inicial = {
    "Insumos Venta": {
        "BOLSA PAPEL CAFE #5 POR PQ/100 PZAS A": 0.832, "BOLSA PAPEL CAFE #6 POR PQ/100 PZAS A": 0.870,
        "BOLSA PAPEL CAFE #14 POR PQ/100 PZAS M": 1.364, "BOLSA PAPEL CAFE #20 POR PQ/100 PZAS M": 1.616,
        "CAJA TUTIS POR PZA A": 0.048, "CAPACILLO CHINO POR PZA B": 0.00104, "CAPACILLO BLANCO POR PZA A": 0.000436,
        "CONT BISAG P/5-6 TUTIS POR PZA A": 0.014, "CUCHARA MED DESCH POR PZA A": 0.00165,
        "EMPLAYE GRANDE ROLLO POR PZA T": 1.174, "PAPEL ALUMINIO POR PZA T": 1.342, "SERVILLETA PQ/500 HJ POR PZA A": 0.001192,
        "BOLSA LOCK POR PZA A": 0.018,
        "AZUCAR REFINADA POR KG A": 1.0, "BOLSA CAMISETA LOGO CH POR KG A": 1.0, 
        "BOLSA CAMISETA LOGO GDE POR KG A": 1.0, "BOLSA NATURAL 18 X 25 POR KG A": 1.0, 
        "PAPEL ENVOLTURA CHAMPLITTE POR KG M": 1.0, "ROLLO POLIPUNTEADO 25 X 35 POR KG B": 1.0, 
        "BOLSA 90 X 120 POR KG A": 1.0, "BOLSA 60 X 90 POR KG M": 1.0
    },
    "Limpieza Venta": {
        "COFIA POR PQ/100 PZAS A": 0.238, "GUANTES TRANSP POLIURETANO POR PQ/100 PZAS A": 0.086,
        "FIBRA PREGON P/BAÑO POR PZA M": 1.0, "FIBRA SCOTCH BRITE POR PZA A": 1.0,
        "FIBRA AZUL P/LAVAR CHAROLAS POR PZA B": 1.0, "CUBETA POR PZA M": 1.0, "ESCOBA POR PZA A": 1.0, 
        "ESCURRIDOR POR PZA M": 1.0, "RECOGEDOR POR PZA M": 1.0, "MECHUDO POR PZA A": 1.0
    },
    "Papelería Venta": {
        "ETIQUETA CHAMPLITTE CHICA 4 X 4 POR PZA B": 0.000328, "ETIQUETA CHAMPLITTE MEDIANA 6 X 6 POR PZA B": 0.00057,
        "GRAPAS CJ POR PZA M": 0.164, "CINTA TRANSP EMPAQUE POR PZA M": 0.272, "CINTA DELIMITADORA POR PZA B": 0.346,
        "COMPROBANTE TRASLADO VALORES POR PZA A": 0.0086, "ETIQUETA BLANCA ADH 13 X 19 POR PQ M": 0.050,
        "HOJAS BLANCAS PQ/500 POR PZA A": 2.146, "TINTA EPSON 544 (CMYK) POR PZA A": 0.078, 
        "ROLLO TERMICO P/TPV POR PZA A": 1.0
    }
}

df_cat_global = conn.query("SELECT * FROM catalogo_productos", ttl="1h")
if df_cat_global.empty:
    with conn.session as s:
        for cat, prods in dicc_inicial.items():
            for art, pue in prods.items():
                s.execute(text("INSERT INTO catalogo_productos (categoria, articulo, pue, ubicacion_conteo, redondeo) VALUES (:c, :a, :p, 'Combinado', 'No') ON CONFLICT DO NOTHING"), 
                          {"c": cat, "a": art, "p": pue})
        s.commit()
    st.cache_data.clear() 
    df_cat_global = conn.query("SELECT * FROM catalogo_productos", ttl="1h")

productos_por_categoria = {}
for _, row in df_cat_global.iterrows():
    c = row['categoria']
    if c not in productos_por_categoria:
        productos_por_categoria[c] = {}
    productos_por_categoria[c][row['articulo']] = row['pue']

for c in ORDEN_CATEGORIAS_OFICIAL:
    if c not in productos_por_categoria:
        productos_por_categoria[c] = {}

# ------------------ SISTEMA DE LOGIN ------------------
def verificar_login():
    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False

    if not st.session_state.autenticado:
        st.markdown("<h2 style='text-align: center;'>⚖️ Baja de insumos</h2>", unsafe_allow_html=True)
        st.markdown("<h4 style='text-align: center; color: gray; margin-bottom: 2rem;'>Control de Acceso</h4>", unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            with st.form("form_login"):
                usuario_input = st.text_input("👤 Usuario:")
                password_input = st.text_input("🔑 Contraseña:", type="password")
                st.write("")
                btn_login = st.form_submit_button("Iniciar Sesión", use_container_width=True, type="primary")
                
                if btn_login:
                    df_check = conn.query("SELECT * FROM usuarios WHERE username = :u AND password = :p", 
                                          params={"u": usuario_input.strip(), "p": password_input}, ttl="10m")
                    if not df_check.empty:
                        st.session_state.autenticado = True
                        st.session_state.usuario_actual = usuario_input.strip()
                        st.session_state.show_toast = "✅ ¡Bienvenid@!"
                        st.rerun()
                    else:
                        st.error("❌ Usuario o contraseña incorrectos.")
        return False
    return True

if not verificar_login():
    st.stop()

# ------------------ BARRA LATERAL ------------------
with st.sidebar:
    st.markdown("### 🏢 Datos de Sesión")
    st.caption(f"👤 Conectado como: **{st.session_state.get('usuario_actual', 'Usuario')}**")
    if st.button("🚪 Cerrar Sesión", use_container_width=True):
        st.session_state.autenticado = False
        st.rerun()
    st.divider()
    
    datos_sucursales = {
        "URANO": "522281342454", "COSTA DE ORO": "522292780850", "COSTA VERDE": "522299359597",
        "DÍAZ MIRÓN": "522291302759", "EJÉRCITO MEXICANO": "522299272107", "PLAZA RÍO": "522299864120",
        "PLAYAS DEL CONCHAL": "522291794020", "COYOL": "522299398334", "LA PLACITA": "522299208481",
        "CUAUHTÉMOC": "522291651340", "MARIO MOLINA": "522291780851", "RAFAEL CUERVO": "522291980229",
        "RÍO MEDIO": "522291005852", "DIVERPLAZA": "522293763180", "BOLÍVAR": "522291002947",
        "CIRCUNVALACIÓN": "522299393726", "J.B. LOBOS": "522299201956", "YÁÑEZ": "522293764940",
        "PALACIO DE HIERRO": "522299272100", "CIUDAD INDUSTRIAL": "522299200278", "DONATO CASAS": "522291653833",
        "LAS VEGAS": "522291932980", "PUENTE MORENO": "522296893999", "CONDESA": "522299863464",
        "MURILLO VIDAL": "522286886443", "ARAUCARIAS": "522281177133", "ÁVILA CAMACHO": "522288170989",
        "EMILIANO ZAPATA": "522969628525"
    }
    
    sucursal_in = st.selectbox("📍 Selecciona tu sucursal:", list(datos_sucursales.keys()))
    numero_wa = datos_sucursales[sucursal_in]
    st.caption(f"📱WhatsApp: **{numero_wa}**")

    st.divider()
    st.markdown("### 🔔 Alertas de Stock")
    umbral_porcentaje = st.number_input("Umbral de Reabastecimiento (%)", min_value=0, value=15, step=1)

    st.divider()
    st.markdown("### 💾 Respaldo de Base de Datos")
    
    with st.form("form_restaurar_boveda"):
        uploaded_csv = st.file_uploader("⬆️ Subir Respaldo CSV", type=["csv"])
        btn_restaurar = st.form_submit_button("🔄 Restaurar Preconteos", use_container_width=True)
        if btn_restaurar:
            if uploaded_csv is not None:
                try:
                    df_upload = pd.read_csv(uploaded_csv)
                    if 'id' in df_upload.columns: df_upload = df_upload.drop(columns=['id'])
                    df_upload['sucursal'] = sucursal_in 
                    df_upload.to_sql("pesajes_guardados", con=conn.engine, if_exists="append", index=False)
                    st.session_state.show_toast = "✅ Respaldo restaurado con éxito"
                    st.rerun()
                except Exception as e: st.error(f"Error al restaurar: {e}")
            else: st.warning("⚠️ Primero selecciona un archivo CSV.")

    st.divider()
    if st.session_state.get('usuario_actual') == 'admin':
        with st.expander("🚨 Zona de Peligro (Formatear Nube)", expanded=False):
            confirmar_borrado = st.checkbox("Confirmar el formateo total")
            if st.button("⚠️ ELIMINAR TODO.", use_container_width=True):
                if not confirmar_borrado: st.error("Debes confirmar primero")
                else:
                    with conn.session as s:
                        s.execute(text("DROP TABLE IF EXISTS pesajes_individuales, pesajes_guardados, auditoria_stock CASCADE"))
                        s.commit()
                    st.session_state.show_toast = "✅ Base de datos eliminada"
                    st.rerun()

# ------------------ ALERTA GLOBAL DE REABASTECIMIENTO ------------------
if "alerta_mostrada" not in st.session_state:
    st.session_state.alerta_mostrada = False

@st.dialog("⚠️ Alerta de Reabastecimiento", width="large")
def dialog_reabastecimiento(df_bajos):
    st.warning(f"Se detectaron {len(df_bajos)} insumos en cero o por debajo del {umbral_porcentaje}%.")
    st.dataframe(df_bajos[['articulo', 'categoria', 'stock', 'pesaje_actual']].rename(
        columns={'articulo': 'Insumo', 'categoria': 'Categoría', 'stock': 'Base Anterior', 'pesaje_actual': 'Pesaje Actual'}
    ), hide_index=True, use_container_width=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Ir a actualizar existencias", type="primary", use_container_width=True):
            st.session_state.alerta_mostrada = True
            st.rerun()
    with col2:
        if st.button("Cerrar", use_container_width=True):
            st.session_state.alerta_mostrada = True
            st.rerun()

# Definir la consulta SIEMPRE para que cualquier pestaña la pueda usar
query_alertas = """
    WITH todos_los_articulos AS (
        SELECT articulo, categoria FROM catalogo_productos
        UNION
        SELECT articulo, categoria FROM auditoria_stock WHERE sucursal = :suc
    )
    SELECT t.articulo, t.categoria,
           COALESCE(a.stock, 0) as stock,
           COALESCE((SELECT SUM(resultado_pue) FROM pesajes_guardados p WHERE p.articulo = t.articulo AND p.sucursal = :suc), 0) +
           COALESCE((SELECT SUM(resultado_pue) FROM pesajes_individuales i WHERE i.articulo = t.articulo AND i.sucursal = :suc), 0) as pesaje_actual
    FROM todos_los_articulos t
    LEFT JOIN auditoria_stock a ON a.articulo = t.articulo AND a.sucursal = :suc
"""

if not st.session_state.alerta_mostrada:
    df_alertas = conn.query(query_alertas, params={"suc": sucursal_in}, ttl=0)
    
    if not df_alertas.empty:
        umbral_decimal = umbral_porcentaje / 100.0
        df_bajos = df_alertas[(df_alertas['pesaje_actual'] <= 0) | (df_alertas['pesaje_actual'] <= (df_alertas['stock'] * umbral_decimal))]
        
        if not df_bajos.empty:
            dialog_reabastecimiento(df_bajos)
        else:
            st.session_state.alerta_mostrada = True
    else:
        st.session_state.alerta_mostrada = True


# ------------------ FUNCIONES AUXILIARES ------------------
def truncar_dos_decimales(valor):
    if valor is None: return 0.0
    s = f"{float(valor):.10f}"
    entero, decimal = s.split('.')
    return float(f"{entero}.{decimal[:2]}")

def formato_estricto(valor):
    if pd.isna(valor) or valor is None: return "0"
    s = f"{float(valor):.10f}" 
    entero, decimal = s.split('.')
    dec_part = decimal[:2]
    return entero if dec_part == "00" else f"{entero}.{dec_part}"

@st.dialog("✅ Registrado", width="large")
def mostrar_popup_exito():
    datos = st.session_state.item_a_guardar
    articulo = datos["articulo"]
    resultado = datos["resultado"]
    sucursal = datos["sucursal"]
    categoria = datos["categoria"]
    peso_bruto = datos["peso_bruto"]
    tara = datos["tara"]
    pue = datos["pue"]
    formula = datos["formula"]
    nuevo_art = datos["nuevo_art"]
    modo_preconteo = datos["modo_preconteo"]
    fecha_mexico = datos["fecha_mexico"]
    num_opciones = datos["num_opciones"]

    st.markdown(f"### 📦 {articulo}")
    
    df_actual_art = conn.query("SELECT * FROM pesajes_individuales WHERE articulo = :art AND sucursal = :suc", params={"art": articulo, "suc": sucursal}, ttl=0)
    df_guardados_art = conn.query("SELECT * FROM pesajes_guardados WHERE articulo = :art AND sucursal = :suc", params={"art": articulo, "suc": sucursal}, ttl=0)
    df_art_combined = pd.concat([df_actual_art, df_guardados_art], ignore_index=True)
    
    sum_anterior = truncar_dos_decimales(df_art_combined['resultado_pue'].sum())
    total_real = truncar_dos_decimales(sum_anterior + resultado)
    
    sumandos = [formato_estricto(val) for val in df_art_combined['resultado_pue']]
    sumandos.append(formato_estricto(resultado))
    texto_total = f"{' + '.join(sumandos)} = {formato_estricto(total_real)}" if len(sumandos) > 1 else formato_estricto(total_real)
    
    st.metric("TOTAL CALCULADO (Sesión + Bóveda)", texto_total)
    st.divider()
    
    df_stock = conn.query("SELECT stock FROM auditoria_stock WHERE articulo = :art AND sucursal = :suc", params={"art": articulo, "suc": sucursal}, ttl=0)
    saved_stock = float(df_stock.iloc[0]['stock']) if not df_stock.empty else None
    
    diferencia_valida = True
    col_st1, col_st2 = st.columns(2)
    with col_st1:
        stock_teorico = st.number_input("Valor en Sistema (Stock):", value=saved_stock, placeholder="Ingresa y presiona Enter", key=f"modal_stock_{articulo}")
        
    with col_st2:
        if stock_teorico is not None:
            diferencia = truncar_dos_decimales(total_real - stock_teorico)
            st.metric("DIFERENCIA", value=" ", delta=formato_estricto(diferencia), delta_color="inverse")
            
            if diferencia > 0:
                diferencia_valida = False
                st.error("⚠️ Diferencia en rojo. El pesaje será omitido (no se guardará).")
    
    st.divider()
    
    col1, col2 = st.columns([1, 1])
    
    def avanzar_y_cerrar():
        if not nuevo_art:
            lista_opciones = datos.get("opciones", [])
            if articulo in lista_opciones:
                idx_actual = lista_opciones.index(articulo)
                if idx_actual < len(lista_opciones) - 1:
                    st.session_state.auto_index = idx_actual + 1
                else:
                    st.session_state.pending_transition = True
            else:
                if st.session_state.auto_index < num_opciones - 1:
                    st.session_state.auto_index += 1
                else:
                    st.session_state.pending_transition = True
        del st.session_state.item_a_guardar

    with col1:
        if st.button("Continuar", type="primary", use_container_width=True):
            if diferencia_valida:
                with conn.session as s:
                    s.execute(text("""INSERT INTO pesajes_individuales 
                                 (sucursal, fecha_hora, articulo, categoria, peso_bruto, tara, pue, resultado_pue, detalle_formula) 
                                 VALUES (:suc, :fh, :art, :cat, :pb, :tara, :pue, :rp, :df)"""),
                              {"suc": sucursal, "fh": fecha_mexico, "art": articulo, "cat": categoria, 
                               "pb": peso_bruto if not modo_preconteo else 0.0, 
                               "tara": tara if not modo_preconteo else 0.0, 
                               "pue": pue if not modo_preconteo else 0.0, 
                               "rp": resultado, "df": formula})
                    if stock_teorico is not None:
                        s.execute(text("""INSERT INTO auditoria_stock (sucursal, articulo, categoria, total_real, stock, diferencia) 
                                     VALUES (:suc, :art, :cat, :tr, :stk, :dif)
                                     ON CONFLICT (sucursal, articulo) DO UPDATE 
                                     SET total_real = EXCLUDED.total_real, stock = EXCLUDED.stock, diferencia = EXCLUDED.diferencia, categoria = EXCLUDED.categoria"""), 
                                  {"suc": sucursal, "art": articulo, "cat": categoria, "tr": total_real, "stk": stock_teorico, "dif": diferencia})
                    s.commit()
            avanzar_y_cerrar()
            st.rerun() 
            
    with col2:
        if diferencia_valida:
            if st.button("📥 Enviar a Bóveda", type="secondary", use_container_width=True):
                with conn.session as s:
                    s.execute(text("""INSERT INTO pesajes_guardados 
                                 (sucursal, fecha_hora, articulo, categoria, peso_bruto, tara, pue, resultado_pue, detalle_formula, aplicado_en_corte) 
                                 VALUES (:suc, :fh, :art, :cat, :pb, :tara, :pue, :rp, :df, FALSE)"""),
                              {"suc": sucursal, "fh": fecha_mexico, "art": articulo, "cat": categoria, 
                               "pb": peso_bruto if not modo_preconteo else 0.0, 
                               "tara": tara if not modo_preconteo else 0.0, 
                               "pue": pue if not modo_preconteo else 0.0, 
                               "rp": resultado, "df": formula})
                    if stock_teorico is not None:
                        s.execute(text("""INSERT INTO auditoria_stock (sucursal, articulo, categoria, total_real, stock, diferencia) 
                                     VALUES (:suc, :art, :cat, :tr, :stk, :dif)
                                     ON CONFLICT (sucursal, articulo) DO UPDATE 
                                     SET total_real = EXCLUDED.total_real, stock = EXCLUDED.stock, diferencia = EXCLUDED.diferencia, categoria = EXCLUDED.categoria"""), 
                                  {"suc": sucursal, "art": articulo, "cat": categoria, "tr": total_real, "stk": stock_teorico, "dif": diferencia})
                    s.commit()
                avanzar_y_cerrar()
                st.session_state.show_toast = "✅ Trasladado a la Bóveda."
                st.rerun()

        if st.button("❌ Cancelar", type="secondary", use_container_width=True):
            del st.session_state.item_a_guardar
            st.rerun()

@st.dialog("⏭️ Confirmar Avance")
def dialog_confirmar_transicion(orden_categorias, orden_ubicaciones, sucursal):
    idx_cat = st.session_state.cat_idx
    idx_ubi = st.session_state.ubi_idx
    
    next_cat_idx = idx_cat
    next_ubi_idx = idx_ubi
    is_complete = False

    if idx_cat < len(orden_categorias) - 1:
        next_cat_idx = idx_cat + 1
    else:
        next_cat_idx = 0
        if idx_ubi < len(orden_ubicaciones) - 1:
            next_ubi_idx = idx_ubi + 1
        else:
            is_complete = True

    if is_complete:
        st.success("🎉 ¡Has completado todas las categorías en todas las ubicaciones!")
        
        if st.button("🔄 Finalizar y Volver al inicio", type="primary", use_container_width=True):
            st.session_state.pending_transition = False
            st.session_state.cat_idx = 0
            st.session_state.ubi_idx = 0
            st.session_state.auto_index = 0
            st.rerun()

    else:
        st.write(f"Terminaste con todos los productos de **{orden_categorias[idx_cat]}**.")
        st.info(f"¿Deseas pasar a **{orden_categorias[next_cat_idx]}** en **{orden_ubicaciones[next_ubi_idx]}**?")
        
        if st.button("✅ Sí, avanzar", type="primary", use_container_width=True):
            st.session_state.cat_idx = next_cat_idx
            st.session_state.ubi_idx = next_ubi_idx
            st.session_state.auto_index = 0
            st.session_state.pending_transition = False
            st.rerun()

        if st.button("❌ No, quedarme aquí", use_container_width=True):
            st.session_state.pending_transition = False
            st.rerun()

def generar_word_tarjetas(df):
    doc = Document()
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21.59), Cm(27.94)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.5)
        
    cols = 3
    rows = (len(df) + cols - 1) // cols or 1
    table = doc.add_table(rows=rows, cols=cols)
    
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_el = OxmlElement(f'w:{edge}')
        border_el.set(qn('w:val'), 'dashed'); border_el.set(qn('w:sz'), '4'); border_el.set(qn('w:space'), '0'); border_el.set(qn('w:color'), '000000')
        tblBorders.append(border_el)
    table._tbl.tblPr.append(tblBorders)
    
    for idx, row_data in df.iterrows():
        r = idx // cols
        c_idx = idx % cols
        cell = table.cell(r, c_idx)
        cell.width = Cm(6)
        table.rows[r].height = Cm(4)
        table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        cell._tc.get_or_add_tcPr().append(tcVAlign)
        
        articulo, resultado_str = str(row_data['articulo']), formato_estricto(row_data['resultado_pue'])
        if "PZA" in articulo.upper() and resultado_str.endswith(".00"): resultado_str = resultado_str[:-3]
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{articulo}\n")
        run1.font.size, run1.bold = Pt(8), True
        run2 = p.add_run(f"\n{resultado_str}")
        run2.font.size, run2.bold = Pt(12), True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------ 4. INTERFAZ PRINCIPAL ------------------
tab_calc, tab_visual, tab_historial, tab_reabasto = st.tabs(["🧮 Pesaje", "🖼️ Esquema Visual", "📋 Reportes", "📦 Reabasto"])

# --- TAB 1: REGISTRO Y AUDITORÍA UNIFICADA (PESAJE) ---
with tab_calc:
    orden_categorias = ORDEN_CATEGORIAS_OFICIAL
    orden_ubicaciones = ["Bodega", "Piso de Venta"]
    
    if "cat_idx" not in st.session_state: st.session_state.cat_idx = 0
    if "ubi_idx" not in st.session_state: st.session_state.ubi_idx = 0
    if "auto_index" not in st.session_state: st.session_state.auto_index = 0
    if "pending_transition" not in st.session_state: st.session_state.pending_transition = False

    if st.session_state.pending_transition:
        dialog_confirmar_transicion(orden_categorias, orden_ubicaciones, sucursal_in)

    modo_seleccionado = st.selectbox("⚙️ Seleccione el Modo:", ["Modo Normal", "Artículo NO listado", "PRE-CONTEO MANUAL (Piezas directas)"], index=0)

    with st.expander("⚙️ Ajustes", expanded=False):
        new_cat = st.selectbox("📂 Seleccione Categoría:", orden_categorias, index=st.session_state.cat_idx)
        new_ubi = st.radio("📍 Ubicación del pesaje:", orden_ubicaciones, index=st.session_state.ubi_idx, horizontal=True)

        if new_cat != orden_categorias[st.session_state.cat_idx]:
            st.session_state.cat_idx = orden_categorias.index(new_cat)
            st.session_state.auto_index = 0
            st.rerun()
        if new_ubi != orden_ubicaciones[st.session_state.ubi_idx]:
            st.session_state.ubi_idx = orden_ubicaciones.index(new_ubi)
            st.session_state.auto_index = 0
            st.rerun()

    categoria_actual = orden_categorias[st.session_state.cat_idx]
    ubicacion_actual = orden_ubicaciones[st.session_state.ubi_idx]
    
    df_cat_global = conn.query("SELECT * FROM catalogo_productos", ttl="1h") 
    df_cat_filtrado = df_cat_global[df_cat_global['categoria'] == categoria_actual]
    
    opciones = []
    for _, row in df_cat_filtrado.iterrows():
        ubi_item = row.get('ubicacion_conteo', 'Combinado')
        if pd.isna(ubi_item) or ubi_item == "":
            ubi_item = "Combinado"
        if ubi_item == "Combinado" or ubi_item.lower() == ubicacion_actual.lower():
            if row['articulo'] not in opciones:
                opciones.append(row['articulo'])
            
    def avanzar_flujo():
        if art_sel in opciones:
            idx_actual = opciones.index(art_sel)
            if idx_actual < len(opciones) - 1:
                st.session_state.auto_index = idx_actual + 1
            else:
                st.session_state.pending_transition = True
        elif len(opciones) > 0 and st.session_state.auto_index < len(opciones) - 1:
            st.session_state.auto_index += 1
        else:
            st.session_state.pending_transition = True

    nuevo_art = (modo_seleccionado == "Artículo NO listado")
    modo_preconteo = (modo_seleccionado == "PRE-CONTEO MANUAL (Piezas directas)")
    
    if not nuevo_art:
        current_index = st.session_state.auto_index
        if current_index >= len(opciones) and len(opciones) > 0: current_index = 0 
        
        art_sel = st.selectbox("📦 Seleccione Artículo:", opciones, index=current_index if len(opciones) > 0 else None, placeholder="Elija un producto...")
        
        df_match = df_cat_filtrado[df_cat_filtrado['articulo'] == art_sel]
        pue_final = float(df_match['pue'].values[0]) if art_sel and not df_match.empty else 1.0
    else:
        c_n1, c_n2 = st.columns([2,1])
        with c_n1: art_sel = st.text_input("Nombre del Nuevo Artículo:", value=None, placeholder="Ej. CAJA PERSONALIZADA")
        with c_n2: pue_final = st.number_input("Asignar Peso Unitario:", value=None, format="%.4f", placeholder="0.0000")

    # ----- POPUP DE PRECONTEO -----
    @st.dialog("⚠️ Preconteo Detectado")
    def dialog_preconteo(articulo, total_preconteo):
        st.info(f"El producto **{articulo}** ya cuenta con un preconteo de **{total_preconteo}**.")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("✅ Omitir y Avanzar", use_container_width=True):
                st.session_state[f"visto_{articulo}"] = True
                avanzar_flujo()
                st.rerun()
        with col2:
            if st.button("➕ Añadir Pesaje", use_container_width=True):
                st.session_state[f"visto_{articulo}"] = True
                st.rerun()
        with col3:
            if st.button("🗑️ Eliminar Preconteo", use_container_width=True):
                with conn.session as s:
                    s.execute(text("DELETE FROM pesajes_guardados WHERE articulo = :art AND sucursal = :suc"), {"art": articulo, "suc": sucursal_in})
                    s.commit()
                st.session_state.show_toast = "✅ Preconteo eliminado."
                st.session_state[f"visto_{articulo}"] = True
                st.rerun()

    if not nuevo_art and art_sel and not st.session_state.get(f"visto_{art_sel}"):
        df_prec = conn.query("SELECT resultado_pue FROM pesajes_guardados WHERE articulo = :art AND sucursal = :suc", params={"art": art_sel, "suc": sucursal_in}, ttl=0)
        if not df_prec.empty:
            dialog_preconteo(art_sel, df_prec['resultado_pue'].sum())
    # ------------------------------

    with st.form(key="form_pesaje", clear_on_submit=True):
        if modo_preconteo:
            st.info("💡 En este modo se registra la cantidad directa sin cálculos de peso.")
            cantidad_directa = st.number_input("Cantidad de piezas (Conteo manual):", value=None, step=1.0, placeholder="Ej. 50")
            peso_bruto, tara_total = 0.0, 0.0
            formula = f"[{ubicacion_actual.upper()}] CONTEO MANUAL DIRECTO"
        else:
            col1, col2 = st.columns([3, 1])
            with col1:
                peso_bruto = st.number_input("⚖️ Peso Bruto (kg):", value=None, format="%.3f", placeholder="0.000")
            with col2:
                st.write("") 
                t_cont = st.checkbox("📦 Tara Contenedor (0.045)", value=False)
                with st.popover("➕ Tara Manual"):
                    t_manual = st.number_input("⚖️ Peso de tara extra:", value=None, format="%.3f", placeholder="0.000")
        
        st.divider()
        
        col_izq, col_centro, col_der = st.columns([1, 2, 1])
        with col_centro:
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                btn_save = st.form_submit_button("📥 GUARDAR Y SIGUIENTE", type="primary", use_container_width=True)
            with col_b2:
                btn_skip = st.form_submit_button("⏭️ OMITIR", use_container_width=True)
                
    if btn_skip:
        if art_sel is not None and art_sel.strip() != "":
            fecha_mexico = datetime.now(zona_mx).strftime("%Y-%m-%d %H:%M:%S")
            try:
                with conn.session as s:
                    s.execute(text("""INSERT INTO pesajes_individuales 
                                 (sucursal, fecha_hora, articulo, categoria, peso_bruto, tara, pue, resultado_pue, detalle_formula) 
                                 VALUES (:suc, :fh, :art, :cat, 0, 0, :pue, 0, '[OMITIDO] Registro en 0')"""),
                              {"suc": sucursal_in, "fh": fecha_mexico, "art": art_sel, "cat": categoria_actual, "pue": pue_final if not modo_preconteo else 0.0})
                    s.commit()
                
                if not nuevo_art:
                    avanzar_flujo()
                    
                st.session_state.show_toast = f"⏭️ {art_sel} omitido"
                st.rerun()
            except Exception as e:
                st.error(f"Error al omitir: {e}")
        else:
            st.error("❌ Selecciona un artículo primero.")

    if btn_save:
        articulo_valido = art_sel is not None and art_sel.strip() != ""
        if modo_preconteo:
            datos_listos = articulo_valido and cantidad_directa is not None
            resultado = truncar_dos_decimales(cantidad_directa) if datos_listos else 0
        else:
            datos_listos = articulo_valido and peso_bruto is not None and pue_final is not None
            if datos_listos:
                tara_total = (0.045 if t_cont else 0) + (t_manual if t_manual is not None else 0.0)
                is_tinta = "TINTA" in str(art_sel).upper(); offset = 0.030 if is_tinta else 0.0
                calc_val = ((peso_bruto - tara_total) - offset) / pue_final
                
                if not nuevo_art:
                    df_match = df_cat_filtrado[df_cat_filtrado['articulo'] == art_sel]
                    red_val = str(df_match['redondeo'].values[0]) if not df_match.empty and 'redondeo' in df_match.columns else "No"
                    
                    if red_val.lower() in ['sí', 'si', 'yes', 'true']:
                        calc_val = float(math.floor(calc_val + 0.5))
                
                resultado = truncar_dos_decimales(calc_val)
                formula = f"[{ubicacion_actual.upper()}] ({peso_bruto:.3f}PB - {tara_total:.3f}T{' - 0.03Env' if is_tinta else ''}) / {pue_final}PUE"

        if datos_listos:
            fecha_mexico = datetime.now(zona_mx).strftime("%Y-%m-%d %H:%M:%S")
            try:
                st.session_state.item_a_guardar = {
                    "articulo": art_sel,
                    "resultado": resultado,
                    "sucursal": sucursal_in,
                    "categoria": categoria_actual,
                    "peso_bruto": peso_bruto if not modo_preconteo else 0.0,
                    "tara": tara_total if not modo_preconteo else 0.0,
                    "pue": pue_final if not modo_preconteo else 0.0,
                    "formula": formula,
                    "nuevo_art": nuevo_art,
                    "modo_preconteo": modo_preconteo,
                    "fecha_mexico": fecha_mexico,
                    "num_opciones": len(opciones),
                    "opciones": opciones
                }
                st.rerun()
            except Exception as e: 
                st.error(f"Error al procesar: {e}")
        else: 
            st.error("❌ Error: Revisa los datos ingresados.")
        
    if "item_a_guardar" in st.session_state:
        mostrar_popup_exito()

    if art_sel:
        st.divider()
        with st.expander(f"📋 Ver detalle e historial de: {art_sel}", expanded=False):
            df_a = conn.query("SELECT * FROM pesajes_individuales WHERE articulo = :art AND sucursal = :suc", params={"art": art_sel, "suc": sucursal_in}, ttl=0)
            df_g = conn.query("SELECT * FROM pesajes_guardados WHERE articulo = :art AND sucursal = :suc", params={"art": art_sel, "suc": sucursal_in}, ttl=0)
            
            if not df_a.empty:
                st.markdown("**🔹 Sesión Actual (Puedes borrar seleccionando la fila y dando clic al basurero)**")
                edited_a = st.data_editor(
                    df_a[['id', 'detalle_formula', 'resultado_pue']].rename(columns={'detalle_formula': 'Operación', 'resultado_pue': 'Cantidad'}), 
                    num_rows="dynamic", hide_index=True, disabled=["Operación", "Cantidad"], key=f"del_a_{art_sel}", use_container_width=True
                )
                ids_del = set(df_a['id']) - set(edited_a['id'])
                if ids_del:
                    if st.button("🗑️ Confirmar Borrado de Pesaje(s)", use_container_width=True):
                        with conn.session as s:
                            for d_id in ids_del: s.execute(text("DELETE FROM pesajes_individuales WHERE id = :id"), {"id": int(d_id)})
                            s.commit()
                        st.session_state.show_toast = "✅ Pesaje eliminado."
                        st.rerun()
            else:
                st.info("No hay pesajes en la sesión actual para este artículo.")

            if not df_g.empty:
                st.markdown("**🔹 Bóveda (Guardados previamente)**")
                df_g['detalle_formula'] = "[GUARDADO] " + df_g['detalle_formula'].astype(str)
                st.dataframe(df_g[['detalle_formula', 'resultado_pue']].rename(columns={'detalle_formula': 'Operación', 'resultado_pue': 'Cantidad'}), hide_index=True, use_container_width=True)


# --- TAB 2: ESQUEMA VISUAL Y STOCK REAL ---
with tab_visual:
    st.subheader("🖼️ Reporte Visual de Insumos")
    fecha_str = datetime.now(zona_mx).strftime("%d/%m/%Y - %H:%M")
    
    html_content = f"""<div style="background-color: white; border-radius: 12px; padding: 20px; box-shadow: 0 4px 10px rgba(0,0,0,0.15); max-width: 900px; margin: auto;">
<div style="text-align: center; color: #8b1c31; font-family: 'Georgia', serif;">
<h1 style="margin: 0; font-size: 32px; font-weight: bold;">Champlitte {sucursal_in.title()}</h1>
<h4 style="margin: 5px 0 15px 0; color: #333; letter-spacing: 2px; font-size: 12px; font-family: sans-serif; font-weight: bold;">CONTROL DE INSUMOS</h4>
<h2 style="margin: 0; font-size: 24px; font-weight: bold;">RESUMEN (TOTALES)</h2>
<p style="color: #666; font-size: 12px; margin-top: 5px; font-family: sans-serif;">{fecha_str}</p>
</div>
<div style="overflow-x: auto;">
<table style="width: 100%; border-collapse: collapse; margin-top: 20px; font-family: sans-serif; font-size: 14px; min-width: 700px;">
<thead>
<tr style="background-color: #8b1c31; color: white; text-align: center; font-size: 12px;">
<th style="padding: 12px; border-top-left-radius: 8px;">CANT. ANTERIOR</th>
<th style="padding: 12px;">CANT. PESADA</th>
<th style="padding: 12px; text-align: left;">PRODUCTO</th>
<th style="padding: 12px;">CANT. A RESTAR</th>
<th style="padding: 12px; border-top-right-radius: 8px;">STOCK ACTUALIZADO</th>
</tr>
</thead>
<tbody>"""
    
    hay_elementos_con_diferencia = False
    row_color_alt = False

    for cat in ORDEN_CATEGORIAS_OFICIAL:
        query_pesajes_raw = '''
            SELECT articulo, resultado_pue 
            FROM (
                SELECT articulo, resultado_pue FROM pesajes_individuales WHERE sucursal = :suc AND categoria = :cat AND detalle_formula NOT LIKE '%[OMITIDO]%'
                UNION ALL
                SELECT articulo, resultado_pue FROM pesajes_guardados WHERE sucursal = :suc AND categoria = :cat AND detalle_formula NOT LIKE '%[OMITIDO]%'
            ) as combinados
        '''
        df_pesajes_raw = conn.query(query_pesajes_raw, params={"suc": sucursal_in, "cat": cat}, ttl=0)
        
        dict_pesajes = {}
        if not df_pesajes_raw.empty:
            for art, group in df_pesajes_raw.groupby('articulo'):
                valores = group['resultado_pue'].tolist()
                total = truncar_dos_decimales(sum(valores))
                str_vals = [formato_estricto(v) for v in valores]
                desglose = f"{' + '.join(str_vals)} = {formato_estricto(total)}" if len(valores) > 1 else formato_estricto(total)
                dict_pesajes[art] = {'total': total, 'desglose': desglose}
                
        df_auditoria = conn.query("SELECT articulo, stock FROM auditoria_stock WHERE sucursal = :suc AND categoria = :cat", params={"suc": sucursal_in, "cat": cat}, ttl=0)
        
        productos_dict = productos_por_categoria.get(cat, {})
        lista_todos = sorted(list(set(list(productos_dict.keys()) + 
                                      (df_auditoria['articulo'].tolist() if not df_auditoria.empty else []) + 
                                      (list(dict_pesajes.keys())))))
        
        filas_categoria = ""
        for art in lista_todos:
            stock_actual = float(df_auditoria[df_auditoria['articulo'] == art]['stock'].iloc[0]) if not df_auditoria.empty and art in df_auditoria['articulo'].values else 0.0
            
            tiene_pesaje = art in dict_pesajes
            
            if tiene_pesaje:
                cant_pesada = dict_pesajes[art]['total']
                str_pesada = dict_pesajes[art]['desglose']
            else:
                cant_pesada = stock_actual 
                str_pesada = formato_estricto(cant_pesada)
                
            cant_a_restar = truncar_dos_decimales(stock_actual - cant_pesada)
            
            if abs(cant_a_restar) < 0.001:
                continue
            
            hay_elementos_con_diferencia = True
            stock_actualizado = cant_pesada  
            
            str_restar = f"-{formato_estricto(abs(cant_a_restar))}" if cant_a_restar > 0 else ""
            str_actual = formato_estricto(stock_actual)
            str_stock_act = formato_estricto(stock_actualizado)
            
            bg_color = "#fffafb" if row_color_alt else "#ffffff"
            row_color_alt = not row_color_alt
            
            filas_categoria += f"""<tr style="background-color: {bg_color}; border-bottom: 1px solid #f0f0f0; text-align: center; color: #8b1c31; font-weight: bold; font-size: 13px;">
<td style="padding: 12px; color: #555; font-weight: normal;">{str_actual}</td>
<td style="padding: 12px;">{str_pesada}</td>
<td style="padding: 12px; text-align: left; color: #333; font-weight: normal;">{art}</td>
<td style="padding: 12px; color: #d9534f; font-weight: bold;">{str_restar}</td>
<td style="padding: 12px; color: #28a745;">{str_stock_act}</td>
</tr>"""

        if filas_categoria:
            html_content += f"""<tr>
<td colspan="5" style="padding: 10px 12px; background-color: #f8eef0; color: #8b1c31; font-weight: bold; text-align: left; font-size: 13px; border-bottom: 2px solid #8b1c31; letter-spacing: 1px;">
📂 {cat.upper()}
</td>
</tr>""" + filas_categoria

    if not hay_elementos_con_diferencia:
        html_content += """<tr>
<td colspan="5" style="padding: 20px; text-align: center; color: #666; font-style: italic;">
No hay diferencias registradas en el stock para esta sucursal.
</td>
</tr>"""

    html_content += """</tbody>
</table>
</div>
</div>"""
    
    st.write(html_content, unsafe_allow_html=True)
    st.divider()

    if st.button("🔄 ACTUALIZAR STOCK PARA MAÑANA (TODAS LAS CATEGORÍAS)", type="primary", use_container_width=True):
        with conn.session as s:
            for cat_upd in ORDEN_CATEGORIAS_OFICIAL:
                query_pesajes_maestro = """
                    SELECT articulo, SUM(resultado_pue) as total_pesado 
                    FROM (
                        SELECT articulo, resultado_pue FROM pesajes_individuales WHERE sucursal = :suc AND categoria = :cat AND detalle_formula NOT LIKE '%[OMITIDO]%'
                        UNION ALL
                        SELECT articulo, resultado_pue FROM pesajes_guardados WHERE sucursal = :suc AND categoria = :cat AND detalle_formula NOT LIKE '%[OMITIDO]%'
                    ) as combinados
                    GROUP BY articulo
                """
                df_pesajes_cat = conn.query(query_pesajes_maestro, params={"suc": sucursal_in, "cat": cat_upd}, ttl=0)
                
                if not df_pesajes_cat.empty:
                    for _, row_p in df_pesajes_cat.iterrows():
                        art_m = row_p["articulo"]
                        nueva_base_m = row_p["total_pesado"]
                        if nueva_base_m > 0:
                            s.execute(text("""INSERT INTO auditoria_stock (sucursal, articulo, categoria, stock, total_real, diferencia) 
                                         VALUES (:suc, :art, :cat, :stk, 0, 0)
                                         ON CONFLICT (sucursal, articulo) DO UPDATE 
                                         SET stock = EXCLUDED.stock"""), 
                                      {"suc": sucursal_in, "art": art_m, "cat": cat_upd, "stk": nueva_base_m})
                
                s.execute(text("DELETE FROM pesajes_individuales WHERE sucursal = :suc AND categoria = :cat"), {"suc": sucursal_in, "cat": cat_upd})
                s.execute(text("DELETE FROM pesajes_guardados WHERE sucursal = :suc AND categoria = :cat"), {"suc": sucursal_in, "cat": cat_upd})
            
            s.commit()
            
        st.session_state.show_toast = "✅ ¡Inventario convertido para mañana en TODAS las categorías!"
        st.rerun()

    st.subheader("📦 Control de Stock Real e Inventario Dinámico")
    st.markdown("Edita directamente la columna **Cantidad Anterior** para calibrar tu base. El sistema restará en automático sumando la sesión normal y la Bóveda.")

    for categoria_activa_stock in ORDEN_CATEGORIAS_OFICIAL:
        st.markdown(f"### 📂 Categoría: {categoria_activa_stock}")
        productos_dict_stock = productos_por_categoria.get(categoria_activa_stock, {})

        query_unificada = """
            SELECT articulo, resultado_pue 
            FROM (
                SELECT articulo, resultado_pue FROM pesajes_individuales WHERE sucursal = :suc AND categoria = :cat AND detalle_formula NOT LIKE '%[OMITIDO]%'
                UNION ALL
                SELECT articulo, resultado_pue FROM pesajes_guardados WHERE sucursal = :suc AND categoria = :cat AND detalle_formula NOT LIKE '%[OMITIDO]%'
            ) as combinados
        """
        df_raw = conn.query(query_unificada, params={"suc": sucursal_in, "cat": categoria_activa_stock}, ttl=0)
        
        pesajes_data = []
        if not df_raw.empty:
            for art, group in df_raw.groupby('articulo'):
                valores = group['resultado_pue'].tolist()
                total = truncar_dos_decimales(sum(valores))
                str_vals = [formato_estricto(v) for v in valores]
                desglose = f"{' + '.join(str_vals)} = {formato_estricto(total)}" if len(valores) > 1 else formato_estricto(total)
                pesajes_data.append({"articulo": art, "total_pesado": total, "desglose_pesada": desglose})
                
        df_total_pesado = pd.DataFrame(pesajes_data) if pesajes_data else pd.DataFrame(columns=["articulo", "total_pesado", "desglose_pesada"])

        df_auditoria_base = conn.query("SELECT articulo, stock FROM auditoria_stock WHERE sucursal = :suc AND categoria = :cat", params={"suc": sucursal_in, "cat": categoria_activa_stock}, ttl=0)

        lista_dict = list(productos_dict_stock.keys())
        lista_audit = df_auditoria_base['articulo'].tolist() if not df_auditoria_base.empty else []
        lista_pesados = df_total_pesado['articulo'].tolist() if not df_total_pesado.empty else []
        lista_todos_articulos = sorted(list(set(lista_dict + lista_audit + lista_pesados)))
        
        df_stock_master = pd.DataFrame({"articulo": lista_todos_articulos})
        
        if not df_auditoria_base.empty:
            df_stock_master = pd.merge(df_stock_master, df_auditoria_base, on="articulo", how="left")
        else:
            df_stock_master["stock"] = 0.0

        if not df_total_pesado.empty:
            df_stock_master = pd.merge(df_stock_master, df_total_pesado, on="articulo", how="left")
        else:
            df_stock_master["total_pesado"] = float('nan') 
            df_stock_master["desglose_pesada"] = "0"

        df_stock_master["stock"] = df_stock_master["stock"].fillna(0.0)
        df_stock_master["total_pesado"] = df_stock_master["total_pesado"].fillna(df_stock_master["stock"])
        df_stock_master["desglose_pesada"] = df_stock_master["desglose_pesada"].fillna("0")
        df_stock_master["cantidad_a_restar"] = df_stock_master["stock"] - df_stock_master["total_pesado"]

        df_stock_display = df_stock_master[[
            "stock", "desglose_pesada", "articulo", "cantidad_a_restar", "total_pesado"
        ]].rename(columns={
            "stock": "Cantidad Anterior",
            "desglose_pesada": "Cantidad Pesada",
            "articulo": "Producto",
            "cantidad_a_restar": "Cantidad a Restar",
            "total_pesado": "Stock Actual"
        })

        with st.expander(f"📦 Tabla de Stock Real - {categoria_activa_stock}", expanded=False):
            df_editado = st.data_editor(
                df_stock_display,
                use_container_width=True,
                hide_index=True,
                column_order=["Cantidad Anterior", "Cantidad Pesada", "Producto", "Cantidad a Restar", "Stock Actual"],
                disabled=["Cantidad Pesada", "Producto", "Cantidad a Restar", "Stock Actual"],
                key=f"editor_stock_real_{categoria_activa_stock}"
            )

            if st.button(f"💾 Guardar Cambios de Stock Inicial ({categoria_activa_stock})", use_container_width=True, key=f"btn_save_stock_{categoria_activa_stock}"):
                with conn.session as s:
                    for _, row in df_editado.iterrows():
                        art = row["Producto"]
                        nuevo_stock = row["Cantidad Anterior"]
                        s.execute(text("""INSERT INTO auditoria_stock (sucursal, articulo, categoria, stock, total_real, diferencia) 
                                     VALUES (:suc, :art, :cat, :stk, 0, 0)
                                     ON CONFLICT (sucursal, articulo) DO UPDATE 
                                     SET stock = EXCLUDED.stock"""), 
                                  {"suc": sucursal_in, "art": art, "cat": categoria_activa_stock, "stk": nuevo_stock})
                    s.commit()
                st.session_state.show_toast = f"✅ Stock inicial de {categoria_activa_stock} actualizado correctamente."
                st.rerun()

        with st.expander(f"📝 Administrar Catálogo: {categoria_activa_stock}", expanded=False):
            st.markdown("Agrega nuevos productos, modifica nombres o cambia el PUE. Al guardar, se aplicará en toda la aplicación sin saltar de pestaña.")
            
            df_cat_global = conn.query("SELECT * FROM catalogo_productos", ttl="1h")
            df_cat_edit = df_cat_global[df_cat_global['categoria'] == categoria_activa_stock][['articulo', 'pue', 'ubicacion_conteo', 'redondeo']].copy()
            
            if 'ubicacion_conteo' not in df_cat_edit.columns: 
                df_cat_edit['ubicacion_conteo'] = "Combinado"
            df_cat_edit['ubicacion_conteo'] = df_cat_edit['ubicacion_conteo'].fillna("Combinado")
            if 'redondeo' not in df_cat_edit.columns:
                df_cat_edit['redondeo'] = "No"
            df_cat_edit['redondeo'] = df_cat_edit['redondeo'].fillna("No")
            
            edited_catalogo = st.data_editor(
                df_cat_edit,
                num_rows="dynamic",
                use_container_width=True,
                hide_index=True,
                key=f"editor_cat_{categoria_activa_stock}",
                column_config={
                    "articulo": st.column_config.TextColumn("Nombre del Producto", required=True),
                    "pue": st.column_config.NumberColumn("Peso Unitario Estándar (PUE)", required=True, format="%.4f"),
                    "ubicacion_conteo": st.column_config.SelectboxColumn(
                        "Aplica en",
                        help="¿Dónde se cuenta este insumo?",
                        options=["Bodega", "Piso de Venta", "Combinado"],
                        required=True,
                        default="Combinado"
                    ),
                    "redondeo": st.column_config.SelectboxColumn(
                        "Redondeo",
                        help="¿Redondear a entero? (.50 sube, < .50 baja)",
                        options=["Sí", "No"],
                        required=True,
                        default="No"
                    )
                }
            )
            
            if st.button(f"💾 Guardar Catálogo ({categoria_activa_stock})", type="secondary", use_container_width=True, key=f"btn_cat_save_{categoria_activa_stock}"):
                with conn.session as s:
                    s.execute(text("DELETE FROM catalogo_productos WHERE categoria = :cat"), {"cat": categoria_activa_stock})
                    
                    for _, row in edited_catalogo.iterrows():
                        art_val = str(row['articulo']).strip()
                        if art_val and not pd.isna(row['pue']):
                            pue_val = float(row['pue'])
                            ubi_val = str(row['ubicacion_conteo']) if pd.notna(row['ubicacion_conteo']) else "Combinado"
                            red_val = str(row['redondeo']) if pd.notna(row['redondeo']) else "No"
                            
                            s.execute(text("INSERT INTO catalogo_productos (categoria, articulo, pue, ubicacion_conteo, redondeo) VALUES (:c, :a, :p, :u, :r) ON CONFLICT DO NOTHING"), 
                                      {"c": categoria_activa_stock, "a": art_val, "p": pue_val, "u": ubi_val, "r": red_val})
                    s.commit()
                st.cache_data.clear() 
                st.toast(f"✅ Catálogo de {categoria_activa_stock} guardado correctamente.")
        st.divider()


# --- TAB 3: EXPORTACIÓN Y BÓVEDA (REPORTES) ---
with tab_historial:
    categorias_disponibles = list(productos_por_categoria.keys())
    categorias_ordenadas = sorted(categorias_disponibles, key=lambda x: ORDEN_CATEGORIAS_OFICIAL.index(x) if x in ORDEN_CATEGORIAS_OFICIAL else 999)
    
    cat_filtro = st.selectbox("📂 Filtrar vistas por Categoría:", ["Todas"] + categorias_ordenadas)
    
    if cat_filtro == "Todas":
        df_actual = conn.query("SELECT * FROM pesajes_individuales WHERE sucursal = :suc", params={"suc": sucursal_in}, ttl=0)
        df_guardados = conn.query("SELECT * FROM pesajes_guardados WHERE sucursal = :suc", params={"suc": sucursal_in}, ttl=0)
    else:
        df_actual = conn.query("SELECT * FROM pesajes_individuales WHERE sucursal = :suc AND categoria = :cat", params={"suc": sucursal_in, "cat": cat_filtro}, ttl=0)
        df_guardados = conn.query("SELECT * FROM pesajes_guardados WHERE sucursal = :suc AND categoria = :cat", params={"suc": sucursal_in, "cat": cat_filtro}, ttl=0)

    df_guardados_rep = df_guardados.copy()
    if not df_guardados_rep.empty: df_guardados_rep['detalle_formula'] = "[GUARDADO] " + df_guardados_rep['detalle_formula'].astype(str)
    df_combined = pd.concat([df_actual, df_guardados_rep], ignore_index=True)

    if not df_combined.empty:
        st.subheader("📄 Tarjetas Recortables (Word)")
        if not df_guardados.empty:
            st.download_button("📄 Descargar Tarjetas en Word (Pre-conteos)", data=generar_word_tarjetas(df_guardados[['articulo', 'resultado_pue']].copy()), file_name=f"Tarjetas_Preconteos_{sucursal_in.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        else: st.info("No hay pre-conteos guardados en la bóveda para generar tarjetas.")
            
        st.markdown(f'<a href="https://wa.me/{numero_wa}" target="_blank" class="btn-wa">💬 ABRIR WHATSAPP (Para enviar archivos)</a>', unsafe_allow_html=True)
        st.divider()
        
        with st.expander("🗑️ Administración de Registros Individuales (Sesión Actual)", expanded=False):
            st.markdown("#### Selecciona el renglón de la izquierda y presiona el ícono de papelera 🗑️ para borrar.")
            edited_df = st.data_editor(df_actual, use_container_width=True, num_rows="dynamic", hide_index=True, disabled=df_actual.columns.tolist(), key="editor_db")
            if st.button("💾 Guardar Cambios en Tabla", use_container_width=True):
                ids_to_delete = set(df_actual['id']) - set(edited_df['id'])
                if ids_to_delete:
                    with conn.session as s:
                        for del_id in ids_to_delete: s.execute(text("DELETE FROM pesajes_individuales WHERE id = :id"), {"id": int(del_id)})
                        s.commit()
                    st.session_state.show_toast = f"✅ Se eliminaron {len(ids_to_delete)} registros correctamente."
                    st.rerun()
                else: st.info("No detecté ninguna fila eliminada para guardar.")

        with st.expander("🛡️ Trasladar a Bóveda (Preconteos Permanentes)", expanded=False):
            st.markdown("Mueve registros de la sesión actual a la bóveda segura.")
            opciones_proteger = df_actual.apply(lambda x: f"ID {x['id']} | {x['articulo']} | {x['resultado_pue']} u.", axis=1).tolist()
            sel = st.multiselect("Selecciona los registros a mover a la bóveda:", opciones_proteger)
            if st.button("📥 Mover seleccionados a la Bóveda") and sel:
                with conn.session as s:
                    for item in sel:
                        id_val = int(item.split(" | ")[0].replace("ID ", ""))
                        s.execute(text("""INSERT INTO pesajes_guardados (sucursal, fecha_hora, articulo, categoria, peso_bruto, tara, pue, resultado_pue, detalle_formula, aplicado_en_corte)
                                     SELECT sucursal, fecha_hora, articulo, categoria, peso_bruto, tara, pue, resultado_pue, detalle_formula, FALSE
                                     FROM pesajes_individuales WHERE id = :id"""), {"id": id_val})
                        s.execute(text("DELETE FROM pesajes_individuales WHERE id = :id"), {"id": id_val})
                    s.commit()
                st.session_state.show_toast = f"✅ Se han trasladado {len(sel)} registros a la Bóveda."
                st.rerun()
            
            st.divider()
            st.markdown("#### 🗃️ Pre-conteos Guardados Actualmente")
            if not df_guardados.empty:
                edited_guardados = st.data_editor(df_guardados, use_container_width=True, num_rows="dynamic", hide_index=True, disabled=df_guardados.columns.tolist(), key="editor_db_guardados")
                if st.button("💾 Eliminar filas borradas de la Bóveda", use_container_width=True):
                    ids_to_delete_g = set(df_guardados['id']) - set(edited_guardados['id'])
                    if ids_to_delete_g:
                        with conn.session as s:
                            for del_id in ids_to_delete_g: s.execute(text("DELETE FROM pesajes_guardados WHERE id = :id"), {"id": int(del_id)})
                            s.commit()
                        st.session_state.show_toast = f"✅ Se eliminaron {len(ids_to_delete_g)} registros guardados."
                        st.rerun()
            else: st.info("No hay pre-conteos guardados en la bóveda.")
    else: st.info(f"No hay pesajes registrados para {sucursal_in} en esta categoría.")

# --- TAB 4: REABASTECIMIENTO ---
with tab_reabasto:
    st.subheader("📦 Lista Crítica y Reabastecimiento")
    st.markdown("Actualiza el stock de los productos detectados en nivel bajo. Al guardar, **las cantidades se trasladarán automáticamente a la base de cálculo de los pesajes**.")
    
    df_estado_actual = conn.query(query_alertas, params={"suc": sucursal_in}, ttl=0)
    
    if not df_estado_actual.empty:
        umbral_decimal = umbral_porcentaje / 100.0
        df_criticos = df_estado_actual[(df_estado_actual['pesaje_actual'] <= 0) | (df_estado_actual['pesaje_actual'] <= (df_estado_actual['stock'] * umbral_decimal))].copy()
        
        if not df_criticos.empty:
            df_criticos['nuevo_stock'] = df_criticos['pesaje_actual']
            
            edited_reabasto = st.data_editor(
                df_criticos[['articulo', 'categoria', 'pesaje_actual', 'nuevo_stock']],
                column_config={
                    "articulo": "Insumo",
                    "categoria": "Categoría",
                    "pesaje_actual": st.column_config.NumberColumn("Stock/Pesaje Actual", disabled=True),
                    "nuevo_stock": st.column_config.NumberColumn("Llegó Reabastecimiento (Ingresa Cantidad)", required=True)
                },
                use_container_width=True, hide_index=True, key="editor_reabasto"
            )
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("💾 Actualizar Existencias en el Sistema", type="primary", use_container_width=True):
                    with conn.session as s:
                        for _, row in edited_reabasto.iterrows():
                            if row['nuevo_stock'] > row['pesaje_actual']:
                                s.execute(text("""
                                    INSERT INTO auditoria_stock (sucursal, articulo, categoria, stock, total_real, diferencia)
                                    VALUES (:suc, :art, :cat, :n_stock, :n_stock, 0)
                                    ON CONFLICT (sucursal, articulo) DO UPDATE 
                                    SET stock = EXCLUDED.stock, total_real = EXCLUDED.total_real, categoria = EXCLUDED.categoria
                                """), {"n_stock": row['nuevo_stock'], "art": row['articulo'], "cat": row['categoria'], "suc": sucursal_in})
                        s.commit()
                    st.session_state.show_toast = "✅ Stock de reabastecimiento trasladado a tablas de pesaje."
                    st.rerun()
            
            with col_btn2:
                lista_texto = "%0A".join([f"- {row['articulo']} (Actual: {row['pesaje_actual']})" for _, row in df_criticos.iterrows()])
                asunto = f"Solicitud Reabastecimiento - {sucursal_in}"
                st.markdown(f'<a href="mailto:compras@ejemplo.com?subject={asunto}&body=Falta stock de los siguientes insumos:%0A{lista_texto}" class="btn-wa" style="background-color:#1f77b4;">✉️ Generar Correo de Pedido</a>', unsafe_allow_html=True)
                
        else:
            st.success("✅ Todo el inventario se encuentra por encima del nivel crítico.")
    else:
        st.info("Aún no hay historial de stock para calcular reabastecimientos.")


components.html("""
    <script>
    const num_inputs = window.parent.document.querySelectorAll('input[type="number"]');
    num_inputs.forEach(input => input.setAttribute('enterkeyhint', 'done'));
    setTimeout(() => {
        const selectores = window.parent.document.querySelectorAll('input[aria-autocomplete="list"], input[role="combobox"]');
        if(selectores.length > 0) selectores[0].focus();
    }, 600); 
    </script>
""", height=0)
